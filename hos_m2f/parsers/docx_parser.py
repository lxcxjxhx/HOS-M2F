"""DOCX解析器"""

import io
from typing import Dict, Any, Optional, Union
from docx import Document

from hos_m2f.model.universal_model import UniversalDocumentModel
from hos_m2f.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):
    """DOCX解析器"""
    
    def parse(self, content: Union[str, bytes], options: Optional[Dict[str, Any]] = None) -> UniversalDocumentModel:
        """解析DOCX内容为中间模型
        
        Args:
            content: DOCX内容
            options: 解析选项
            
        Returns:
            UniversalDocumentModel: 通用文档模型
        """
        options = options or {}
        
        # 创建DOCX文档对象
        if isinstance(content, bytes):
            doc = Document(io.BytesIO(content))
        else:
            doc = Document(content)
        
        # 创建通用文档模型
        model = UniversalDocumentModel()
        
        # 1. 提取元数据
        meta = self._extract_meta(doc)
        model.set_meta(meta)
        
        # 2. 生成HTML内容
        html_content = self._generate_html_content(doc, options)
        model.set_html_content(html_content)
        
        # 3. 分析文档结构
        structure = self._analyze_structure(doc)
        for item in structure:
            model.add_structure_item(item)
        
        # 4. 提取语义信息
        semantics = self._extract_semantics(doc)
        model.get_json()["semantics"].update(semantics)
        
        return model
    
    def _extract_meta(self, doc: Document) -> Dict[str, Any]:
        """提取元数据
        
        Args:
            doc: DOCX文档对象
            
        Returns:
            Dict[str, Any]: 元数据字典
        """
        meta = {}
        
        # 尝试从DOCX属性中提取元数据
        try:
            core_properties = doc.core_properties
            if core_properties.title:
                meta["title"] = core_properties.title
            if core_properties.author:
                meta["author"] = core_properties.author
            if core_properties.subject:
                meta["description"] = core_properties.subject
            if core_properties.keywords:
                meta["tags"] = [tag.strip() for tag in core_properties.keywords.split(';')]
            if core_properties.created:
                meta["publish_date"] = core_properties.created.strftime('%Y-%m-%d')
        except Exception as e:
            print(f"Warning: Failed to extract core properties: {e}")
        
        # 如果没有标题，尝试从第一个段落提取
        if not meta.get("title") and len(doc.paragraphs) > 0:
            first_paragraph = doc.paragraphs[0].text.strip()
            if first_paragraph:
                meta["title"] = first_paragraph
        
        return meta
    
    def _generate_html_content(self, doc: Document, options: Dict[str, Any]) -> str:
        """生成HTML内容
        
        Args:
            doc: DOCX文档对象
            options: 生成选项
            
        Returns:
            str: HTML内容
        """
        html_parts = []
        
        # 处理段落
        for paragraph in doc.paragraphs:
            html = self._render_paragraph(paragraph)
            if html:
                html_parts.append(html)
        
        # 处理表格
        for table in doc.tables:
            html = self._render_table(table)
            if html:
                html_parts.append(html)
        
        return "<div class='content'>" + "\n".join(html_parts) + "</div>"
    
    def _render_paragraph(self, paragraph: Any) -> str:
        """渲染段落
        
        Args:
            paragraph: 段落对象
            
        Returns:
            str: HTML内容
        """
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # 检测标题级别
        style_name = paragraph.style.name
        if style_name.startswith('Heading '):
            level = int(style_name.split(' ')[1])
            return f"<h{level}>{text}</h{level}>"
        
        # 检测列表
        if paragraph.style.name.startswith('List '):
            if paragraph.style.name.startswith('List Bullet'):
                return f"<li>{text}</li>"
            elif paragraph.style.name.startswith('List Number'):
                return f"<li>{text}</li>"
        
        # 默认段落
        return f"<p>{text}</p>"
    
    def _render_table(self, table: Any) -> str:
        """渲染表格
        
        Args:
            table: 表格对象
            
        Returns:
            str: HTML内容
        """
        html = "<table>"
        
        # 处理表头
        if table.rows:
            html += "<thead><tr>"
            for cell in table.rows[0].cells:
                html += f"<th>{cell.text.strip()}</th>"
            html += "</tr></thead>"
        
        # 处理表格内容
        html += "<tbody>"
        for i, row in enumerate(table.rows):
            if i == 0:  # 跳过表头
                continue
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{cell.text.strip()}</td>"
            html += "</tr>"
        html += "</tbody>"
        
        html += "</table>"
        return html
    
    def _analyze_structure(self, doc: Document) -> list:
        """分析文档结构
        
        Args:
            doc: DOCX文档对象
            
        Returns:
            list: 结构项列表
        """
        structure = []
        
        # 分析段落
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            style_name = paragraph.style.name
            if style_name.startswith('Heading '):
                level = int(style_name.split(' ')[1])
                structure.append({
                    "type": "heading",
                    "level": level,
                    "title": text,
                    "position": i
                })
            elif paragraph.style.name.startswith('List '):
                list_type = "ordered" if paragraph.style.name.startswith('List Number') else "unordered"
                if not any(item.get('type') == 'list' and item.get('position') == i for item in structure):
                    structure.append({
                        "type": "list",
                        "list_type": list_type,
                        "position": i
                    })
        
        # 分析表格
        for i, table in enumerate(doc.tables):
            structure.append({
                "type": "table",
                "position": i
            })
        
        return structure
    
    def _extract_semantics(self, doc: Document) -> Dict[str, Any]:
        """提取语义信息
        
        Args:
            doc: DOCX文档对象
            
        Returns:
            Dict[str, Any]: 语义信息字典
        """
        semantics = {
            "domain_tag": "",
            "section_id": {},
            "table_type": {},
            "list_type": {}
        }
        
        # 提取章节ID
        section_counter = 1
        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name
            if style_name.startswith('Heading '):
                level = int(style_name.split(' ')[1])
                text = paragraph.text.strip()
                section_id = f"sec-{section_counter}"
                semantics["section_id"][section_id] = {
                    "title": text,
                    "level": level,
                    "position": i
                }
                section_counter += 1
        
        # 提取列表类型
        list_counter = 0
        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name
            if style_name.startswith('List '):
                list_type = "ordered" if style_name.startswith('List Number') else "unordered"
                if not any(item.get('type') == 'list' and item.get('position') == i for item in semantics["list_type"]):
                    semantics["list_type"][f"list-{list_counter}"] = list_type
                    list_counter += 1
        
        return semantics
