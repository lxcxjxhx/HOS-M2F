"""HTML到DOCX格式转换器"""

from typing import Any, Optional, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from hos_m2f.converters.base_converter import BaseConverter
import re
from bs4 import BeautifulSoup


class HTMLToDOCXConverter(BaseConverter):
    """HTML到DOCX格式转换器"""
    
    def convert(self, input_content: str, options: Optional[Dict[str, Any]] = None) -> bytes:
        """将HTML转换为DOCX
        
        Args:
            input_content: HTML内容
            options: 转换选项
                - template: 企业模板文件路径
                - toc: 是否生成目录
                - number_sections: 是否添加章节编号
            
        Returns:
            bytes: DOCX文件的二进制数据
        """
        if options is None:
            options = {}
        
        import os
        
        # 获取模板路径
        template = options.get('template', None)
        
        # 创建文档（使用模板或默认）
        if template and os.path.exists(template):
            doc = Document(template)
        else:
            # 尝试使用默认模板
            default_template = os.path.join(os.path.dirname(__file__), '..', 'templates', 'reference.docx')
            if os.path.exists(default_template):
                doc = Document(default_template)
            else:
                doc = Document()
        
        # 设置默认样式
        self._setup_styles(doc)
        
        # 解析HTML
        soup = BeautifulSoup(input_content, 'html.parser')
        
        # 处理HTML内容
        self._process_html_content(soup, doc)
        
        # 检查是否需要生成目录
        if options.get('toc', False):
            # 在文档开头添加目录
            doc.add_page_break()
            doc.add_heading('目录', level=0)
            doc.add_paragraph().add_run().add_field('TOC')
            doc.add_page_break()
        
        # 保存为二进制数据
        import io
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()
    
    def _setup_styles(self, doc):
        """设置文档样式"""
        styles = doc.styles
        
        # 设置正文样式
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本
        
        # 设置段落格式
        paragraph_format = normal_style.paragraph_format
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.space_after = Pt(12)  # 段后12磅
        
        # 设置标题样式
        for i in range(1, 6):
            heading_style = styles[f'Heading {i}']
            font = heading_style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(14 + (6 - i) * 2)
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本
            
            # 设置标题段落格式
            heading_paragraph_format = heading_style.paragraph_format
            heading_paragraph_format.line_spacing = 1.2  # 1.2倍行距
            heading_paragraph_format.space_after = Pt(12)  # 段后12磅
            heading_paragraph_format.space_before = Pt(0)  # 段前0磅
        
        # 设置列表样式
        for style_name in ['List Bullet', 'List Number']:
            if style_name in styles:
                list_style = styles[style_name]
                font = list_style.font
                font.name = 'Microsoft YaHei'
                font.size = Pt(12)
                font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本
                
                # 设置列表段落格式
                list_paragraph_format = list_style.paragraph_format
                list_paragraph_format.line_spacing = 1.5  # 1.5倍行距
                list_paragraph_format.space_after = Pt(6)  # 段后6磅
    
    def _process_html_content(self, soup, doc):
        """处理HTML内容"""
        # 处理<body>标签内的内容
        body = soup.find('body')
        if body:
            self._process_element(body, doc)
        else:
            # 如果没有<body>标签，处理整个文档
            self._process_element(soup, doc)
    
    def _process_element(self, element, doc, parent=None, list_level=0):
        """处理HTML元素"""
        if element.name == 'h1':
            # 处理一级标题
            text = element.get_text(strip=True)
            doc.add_heading(text, level=0)
        elif element.name in ['h2', 'h3', 'h4', 'h5', 'h6']:
            # 处理其他级别的标题
            level = int(element.name[1]) - 1
            text = element.get_text(strip=True)
            doc.add_heading(text, level=level)
        elif element.name == 'p':
            # 处理段落
            text = element.get_text(strip=False)
            if text.strip():
                p = doc.add_paragraph()
                self._process_inline_elements(element, p)
        elif element.name == 'div':
            # 处理div标签
            for child in element.children:
                self._process_element(child, doc, parent)
        elif element.name in ['ul', 'ol']:
            # 处理列表
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph()
                # 设置列表样式
                if element.name == 'ul':
                    p.style = 'List Bullet'
                else:
                    p.style = 'List Number'
                # 缩进
                p.paragraph_format.left_indent = Inches(0.5 * list_level)
                # 处理列表项内容
                self._process_inline_elements(li, p)
                # 处理嵌套列表
                for child in li.children:
                    if child.name in ['ul', 'ol']:
                        self._process_element(child, doc, p, list_level + 1)
        elif element.name == 'table':
            # 处理表格
            self._process_table(element, doc)
        elif element.name == 'img':
            # 处理图片
            self._process_image(element, doc)
        elif element.name == 'a':
            # 处理链接
            if parent:
                text = element.get_text(strip=True)
                href = element.get('href', '')
                if text and href:
                    run = parent.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    run.underline = WD_UNDERLINE.SINGLE
                    parent.add_run(f' ({href})')
        elif element.name == 'strong' or element.name == 'b':
            # 处理粗体
            if parent:
                text = element.get_text(strip=False)
                run = parent.add_run(text)
                run.bold = True
        elif element.name == 'em' or element.name == 'i':
            # 处理斜体
            if parent:
                text = element.get_text(strip=False)
                run = parent.add_run(text)
                run.italic = True
        elif element.name == 'code':
            # 处理代码
            if parent:
                text = element.get_text(strip=False)
                run = parent.add_run(text)
                run.font.name = 'Courier New'
        elif element.name == 'pre':
            # 处理代码块
            text = element.get_text(strip=False)
            p = doc.add_paragraph('代码:')
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        elif element.name == 'br':
            # 处理换行
            if parent:
                parent.add_run('\n')
        elif isinstance(element, str):
            # 处理文本节点
            if parent and element.strip():
                parent.add_run(element)
    
    def _process_inline_elements(self, element, parent):
        """处理内联元素"""
        for child in element.children:
            if isinstance(child, str):
                # 处理文本节点
                if child.strip():
                    parent.add_run(child)
            else:
                # 处理内联元素
                if child.name == 'strong' or child.name == 'b':
                    text = child.get_text(strip=False)
                    run = parent.add_run(text)
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    text = child.get_text(strip=False)
                    run = parent.add_run(text)
                    run.italic = True
                elif child.name == 'code':
                    text = child.get_text(strip=False)
                    run = parent.add_run(text)
                    run.font.name = 'Courier New'
                elif child.name == 'a':
                    text = child.get_text(strip=True)
                    href = child.get('href', '')
                    if text and href:
                        run = parent.add_run(text)
                        run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                        run.underline = WD_UNDERLINE.SINGLE
                        parent.add_run(f' ({href})')
                elif child.name == 'br':
                    parent.add_run('\n')
                else:
                    # 递归处理其他内联元素
                    self._process_inline_elements(child, parent)
    
    def _process_table(self, table, doc):
        """处理表格"""
        try:
            # 获取所有行
            rows = table.find_all('tr')
            if not rows:
                return
            
            # 获取表头
            header_row = rows[0]
            header_cells = header_row.find_all(['th', 'td'])
            if not header_cells:
                return
            
            # 创建表格
            doc_table = doc.add_table(rows=1, cols=len(header_cells))
            doc_table.style = 'Table Grid'
            
            # 填充表头
            for i, cell in enumerate(header_cells):
                text = cell.get_text(strip=True)
                doc_cell = doc_table.rows[0].cells[i]
                p = doc_cell.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 填充数据行
            for row in rows[1:]:
                cells = row.find_all(['th', 'td'])
                if cells:
                    new_row = doc_table.add_row()
                    for i, cell in enumerate(cells):
                        if i < len(new_row.cells):
                            text = cell.get_text(strip=True)
                            doc_cell = new_row.cells[i]
                            p = doc_cell.add_paragraph(text)
        except Exception as e:
            # 如果处理失败，添加表格描述
            doc.add_paragraph(f'Table error: {str(e)}')
    
    def _process_image(self, img, doc):
        """处理图片"""
        try:
            src = img.get('src', '')
            alt = img.get('alt', '')
            title = img.get('title', '')
            
            # 使用alt或title作为替代文本
            text = alt or title or 'Image'
            
            if src:
                import os
                import requests
                from io import BytesIO
                
                # 检查是否是本地图片
                if os.path.exists(src):
                    # 添加本地图片
                    doc.add_picture(src)
                else:
                    # 尝试从网络获取图片
                    try:
                        response = requests.get(src, timeout=5)
                        if response.status_code == 200:
                            # 添加远程图片
                            image_stream = BytesIO(response.content)
                            doc.add_picture(image_stream)
                        else:
                            # 如果获取失败，添加图片描述
                            doc.add_paragraph(f'Image: {text} ({src})')
                    except:
                        # 如果获取失败，添加图片描述
                        doc.add_paragraph(f'Image: {text} ({src})')
            else:
                # 如果没有src属性，添加图片描述
                doc.add_paragraph(f'Image: {text}')
        except Exception as e:
            # 如果处理失败，添加图片描述
            doc.add_paragraph(f'Image error: {str(e)}')
    
    def get_supported_formats(self) -> tuple:
        """获取支持的格式"""
        return ('html', 'docx')