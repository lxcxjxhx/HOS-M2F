"""DOCX到Markdown格式转换器"""

from typing import Any, Optional, Dict, List, Tuple
from hos_m2f.converters.base_converter import BaseConverter
from docx import Document
import os
import io


class DOCXToMDConverter(BaseConverter):
    """DOCX到Markdown格式转换器"""
    
    def convert(self, input_content: bytes, options: Optional[Dict[str, Any]] = None) -> bytes:
        """将DOCX转换为Markdown
        
        Args:
            input_content: DOCX文件的二进制数据
            options: 转换选项
                - images_dir: 图片保存目录
                - include_images: 是否包含图片
                - mermaid_detection: 是否尝试检测Mermaid图表
            
        Returns:
            bytes: Markdown文件的二进制数据
        """
        if options is None:
            options = {}
        
        # 加载DOCX文档
        doc = Document(io.BytesIO(input_content))
        
        # 转换为Markdown
        md_content = []
        
        # 处理图片
        images_dir = options.get('images_dir', 'images')
        include_images = options.get('include_images', True)
        mermaid_detection = options.get('mermaid_detection', True)
        
        extracted_images = []
        if include_images:
            # 创建images目录
            os.makedirs(images_dir, exist_ok=True)
            # 提取图片
            extracted_images = self._extract_images(doc, images_dir)
        
        # 处理段落
        in_mermaid_block = False
        mermaid_code = []
        
        for paragraph in doc.paragraphs:
            # 检测Mermaid图表开始
            if not in_mermaid_block and self._is_mermaid_start(paragraph.text):
                in_mermaid_block = True
                # 移除"Mermaid Chart:"这样的前缀
                text = paragraph.text.strip()
                if text.startswith('Mermaid Chart:'):
                    text = text[len('Mermaid Chart:'):].strip()
                mermaid_code = [text] if text else []
                continue
            
            # 检测Mermaid图表结束
            if in_mermaid_block:
                text = paragraph.text.strip()
                
                # 检查是否是章节标题或其他非Mermaid内容
                if self._is_mermaid_end(text):
                    # 处理Mermaid图表
                    if mermaid_code:
                        md_content.append('```mermaid')
                        for line in mermaid_code:
                            if line.strip():
                                md_content.append(line)
                        md_content.append('```')
                        md_content.append('')
                    in_mermaid_block = False
                    mermaid_code = []
                else:
                    # 只添加非空行
                    if text:
                        mermaid_code.append(text)
                continue
            
            # 处理标题
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split(' ')[1])
                md_content.append('#' * level + ' ' + paragraph.text)
            # 处理列表
            elif paragraph.style.name in ['List Bullet', 'List Number']:
                # 检测缩进级别
                try:
                    if paragraph.paragraph_format.left_indent:
                        indent_level = int(paragraph.paragraph_format.left_indent.inches // 0.5)
                    else:
                        indent_level = 0
                except AttributeError:
                    indent_level = 0
                
                prefix = '  ' * indent_level
                
                if paragraph.style.name == 'List Number':
                    # 简化处理，实际项目中需要更复杂的列表编号处理
                    md_content.append(prefix + '1. ' + paragraph.text)
                else:
                    md_content.append(prefix + '- ' + paragraph.text)
            # 处理普通段落
            else:
                if paragraph.text:
                    # 处理文本格式
                    formatted_text = self._process_run_formatting(paragraph)
                    md_content.append(formatted_text)
            
            # 添加空行
            md_content.append('')
        
        # 处理表格
        for table in doc.tables:
            # 表头
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text)
            
            # 表格分隔线
            separators = ['---'] * len(headers)
            
            # 表格数据
            rows = []
            for row in table.rows[1:]:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text)
                rows.append(row_cells)
            
            # 生成Markdown表格
            if headers:
                md_content.append('| ' + ' | '.join(headers) + ' |')
                md_content.append('| ' + ' | '.join(separators) + ' |')
                for row in rows:
                    md_content.append('| ' + ' | '.join(row) + ' |')
                md_content.append('')
        
        # 添加图片
        if extracted_images:
            md_content.append('## 图片')
            md_content.append('')
            for image in extracted_images:
                image_path = os.path.join(images_dir, image['name'])
                # 检查是否可能是Mermaid图表
                is_mermaid = False
                if mermaid_detection:
                    is_mermaid = self._is_mermaid_image(image)
                
                if is_mermaid:
                    md_content.append(f'### Mermaid Chart')
                    md_content.append(f'![Mermaid Chart]({image_path})')
                    md_content.append('')
                    md_content.append('```mermaid')
                    md_content.append('# Mermaid chart image - original code not available')
                    md_content.append('# You can edit this section to add the original Mermaid code')
                    md_content.append('```')
                else:
                    md_content.append(f'![Image]({image_path})')
                md_content.append('')
        
        # 生成Markdown内容
        md_text = '\n'.join(md_content)
        
        return md_text.encode('utf-8')
    
    def split_docx(self, input_content: bytes, output_dir: str, options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """拆分DOCX文件为LaTeX格式和图片
        
        Args:
            input_content: DOCX文件的二进制数据
            output_dir: 输出目录
            options: 拆分选项
            
        Returns:
            Dict[str, Any]: 拆分结果
        """
        if options is None:
            options = {}
        
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 创建子目录
        images_dir = os.path.join(output_dir, 'images')
        os.makedirs(images_dir, exist_ok=True)
        
        # 加载DOCX文档
        doc = Document(io.BytesIO(input_content))
        
        # 提取图片
        images = self._extract_images(doc, images_dir)
        
        # 生成LaTeX格式
        latex_content = self._generate_latex(doc, images, options)
        
        # 保存LaTeX文件
        latex_path = os.path.join(output_dir, 'document.tex')
        with open(latex_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        return {
            'success': True,
            'latex_path': latex_path,
            'images': images,
            'images_dir': images_dir
        }
    
    def _extract_images(self, doc: Document, images_dir: str) -> List[Dict[str, Any]]:
        """提取DOCX中的图片
        
        Args:
            doc: DOCX文档对象
            images_dir: 图片输出目录
            
        Returns:
            List[Dict[str, Any]]: 提取的图片信息
        """
        images = []
        
        # 提取文档中的图片
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                # 获取图片内容
                image_part = rel.target_part
                image_content = image_part.blob
                
                # 确定图片扩展名
                content_type = image_part.content_type
                if content_type == 'image/jpeg':
                    ext = '.jpg'
                elif content_type == 'image/png':
                    ext = '.png'
                elif content_type == 'image/gif':
                    ext = '.gif'
                else:
                    ext = '.bin'
                
                # 保存图片
                image_name = f'image_{i+1}{ext}'
                image_path = os.path.join(images_dir, image_name)
                with open(image_path, 'wb') as f:
                    f.write(image_content)
                
                # 记录图片信息
                images.append({
                    'name': image_name,
                    'path': image_path,
                    'content_type': content_type,
                    'content': image_content
                })
        
        return images
    
    def _is_mermaid_image(self, image: Dict[str, Any]) -> bool:
        """检测图片是否可能是Mermaid图表
        
        Args:
            image: 图片信息字典
            
        Returns:
            bool: 是否可能是Mermaid图表
        """
        # 基于图片特征检测Mermaid图表
        # 1. 检查文件名是否包含mermaid相关关键词
        if 'mermaid' in image['name'].lower():
            return True
        
        # 2. 检查图片内容大小和类型
        # Mermaid图表通常是PNG格式，且大小相对较小
        if image['content_type'] == 'image/png':
            # 检查图片大小（示例值，可根据实际情况调整）
            content_size = len(image.get('content', ''))
            if 1000 < content_size < 1000000:  # 1KB - 1MB
                return True
        
        # 3. 检查图片内容（如果需要更复杂的检测）
        # 这里可以添加更复杂的图像分析逻辑
        # 例如使用图像处理库分析颜色分布、图案等
        
        return False
    
    def _is_mermaid_start(self, text: str) -> bool:
        """检测Mermaid图表开始
        
        Args:
            text: 段落文本
            
        Returns:
            bool: 是否是Mermaid图表开始
        """
        # 检测常见的Mermaid图表开始标记
        mermaid_start_patterns = [
            'Mermaid Chart:',
            'graph TD',
            'graph LR',
            'graph TB',
            'graph BT',
            'flowchart TD',
            'flowchart LR',
            'flowchart TB',
            'flowchart BT',
            'sequenceDiagram',
            'classDiagram',
            'stateDiagram',
            'entityRelationshipDiagram',
            'userJourney',
            'gantt',
            'pie',
            'gitGraph',
            'requirementDiagram'
        ]
        
        # 添加中文前缀检测
        chinese_prefixes = [
            '架构图',
            '流程图',
            '关系图',
            '状态图',
            '时序图',
            '类图',
            '用户旅程图',
            '甘特图',
            '饼图',
            'Git图',
            '需求图'
        ]
        
        text = text.strip()
        
        # 检查是否包含英文Mermaid标记
        if any(pattern in text for pattern in mermaid_start_patterns):
            return True
        
        # 检查是否包含中文前缀
        for prefix in chinese_prefixes:
            if prefix in text:
                # 提取前缀后面的内容，检查是否包含Mermaid代码
                suffix = text.split(prefix)[-1].strip()
                # 移除可能的冒号
                if suffix.startswith('：') or suffix.startswith(':'):
                    suffix = suffix[1:].strip()
                # 检查后缀是否包含Mermaid代码标记
                if any(pattern in suffix for pattern in mermaid_start_patterns if pattern != 'Mermaid Chart:'):
                    return True
        
        return False
    
    def _is_mermaid_end(self, text: str) -> bool:
        """检测Mermaid图表结束
        
        Args:
            text: 段落文本
            
        Returns:
            bool: 是否是Mermaid图表结束
        """
        # 检测Mermaid图表结束
        text = text.strip()
        
        # 1. 空行
        if not text:
            return True
        
        # 2. 下一个标题开始
        if text.startswith('#'):
            return True
        
        # 3. 列表开始
        if text.startswith('- ') or text.startswith('1. '):
            return True
        
        # 4. 表格开始
        if text.startswith('|'):
            return True
        
        # 5. 代码块开始
        if text.startswith('```'):
            return True
        
        # 6. 下一个章节标题（如"3.4.2"这样的格式）
        if any(char.isdigit() for char in text) and '.' in text:
            # 检查是否是章节编号格式
            parts = text.split('.')
            if len(parts) >= 2:
                # 提取数字部分
                num_parts = []
                for part in parts:
                    num_part = ''.join(c for c in part if c.isdigit())
                    if num_part:
                        num_parts.append(num_part)
                if len(num_parts) >= 2:
                    return True
        
        # 7. 检查是否是Mermaid代码的一部分
        # Mermaid代码通常包含特定的关键词和语法
        mermaid_keywords = [
            'graph', 'flowchart', 'sequenceDiagram', 'classDiagram',
            'stateDiagram', 'entityRelationshipDiagram', 'userJourney',
            'gantt', 'pie', 'gitGraph', 'requirementDiagram',
            'style', 'click', 'link', 'classDef', 'state', 'participant',
            'as', '->', '-->', '==>', 'note', 'opt', 'alt', 'par', 'loop'
        ]
        
        # 如果文本不包含任何Mermaid关键词，可能是图表结束
        if not any(keyword in text for keyword in mermaid_keywords):
            return True
        
        return False
    
    def _process_run_formatting(self, paragraph) -> str:
        """处理段落中的文本格式
        
        Args:
            paragraph: 段落对象
            
        Returns:
            str: 格式化后的文本
        """
        formatted_text = []
        
        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue
            
            # 处理粗体
            if run.bold:
                run_text = f'**{run_text}**'
            
            # 处理斜体
            if run.italic:
                run_text = f'*{run_text}*'
            
            # 处理下划线
            if run.underline:
                run_text = f'<u>{run_text}</u>'
            
            # 处理字体颜色
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                # 正确处理RGBColor对象
                try:
                    # 尝试直接访问属性
                    hex_color = f'#{rgb.red:02x}{rgb.green:02x}{rgb.blue:02x}'
                except AttributeError:
                    # 尝试其他可能的属性名
                    try:
                        hex_color = f'#{rgb.r:02x}{rgb.g:02x}{rgb.b:02x}'
                    except AttributeError:
                        # 如果都失败，跳过颜色处理
                        hex_color = None
                
                if hex_color:
                    run_text = f'<span style="color:{hex_color}">{run_text}</span>'
            
            # 处理字体大小
            if run.font.size:
                size = run.font.size.pt
                run_text = f'<span style="font-size:{size}pt">{run_text}</span>'
            
            # 处理等宽字体（可能是代码）
            if run.font.name and 'courier' in run.font.name.lower():
                run_text = f'`{run_text}`'
            
            formatted_text.append(run_text)
        
        return ''.join(formatted_text)
    
    def _generate_latex(self, doc: Document, images: List[Dict[str, Any]], options: Dict[str, Any]) -> str:
        """生成LaTeX格式
        
        Args:
            doc: DOCX文档对象
            images: 提取的图片信息
            options: 生成选项
            
        Returns:
            str: LaTeX内容
        """
        latex_content = []
        
        # 添加LaTeX文档头部
        latex_content.append('\\documentclass{article}')
        latex_content.append('\\usepackage{ctex}')
        latex_content.append('\\usepackage{graphicx}')
        latex_content.append('\\usepackage{hyperref}')
        latex_content.append('\\usepackage{booktabs}')
        latex_content.append('\\begin{document}')
        
        # 处理段落
        image_index = 0
        for paragraph in doc.paragraphs:
            # 处理标题
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split(' ')[1])
                if level == 1:
                    latex_content.append(f'\\section{{{paragraph.text}}}')
                elif level == 2:
                    latex_content.append(f'\\subsection{{{paragraph.text}}}')
                elif level == 3:
                    latex_content.append(f'\\subsubsection{{{paragraph.text}}}')
                else:
                    latex_content.append(f'\\paragraph{{{paragraph.text}}}')
            # 处理列表
            elif paragraph.style.name in ['List Bullet', 'List Number']:
                # 检测缩进级别
                indent_level = int(paragraph.paragraph_format.left_indent.inches // 0.5)
                
                if indent_level == 0:
                    latex_content.append('\\begin{itemize}')
                    latex_content.append(f'\\item {paragraph.text}')
                    latex_content.append('\\end{itemize}')
                else:
                    # 简化处理，实际项目中需要更复杂的列表处理
                    latex_content.append(f'\\item {paragraph.text}')
            # 处理普通段落
            else:
                if paragraph.text:
                    latex_content.append(paragraph.text)
            
            # 添加空行
            latex_content.append('')
        
        # 处理表格
        for table in doc.tables:
            # 表头
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text)
            
            # 表格数据
            rows = []
            for row in table.rows[1:]:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text)
                rows.append(row_cells)
            
            # 生成LaTeX表格
            if headers:
                latex_content.append('\\begin{table}[htbp]')
                latex_content.append('\\centering')
                latex_content.append(f'\\begin{{tabular}}{{{"|l" * len(headers)}|}}')
                latex_content.append('\\hline')
                
                # 添加表头
                header_line = ' & '.join(headers)
                latex_content.append(f'{header_line} \\')
                latex_content.append('\\hline')
                
                # 添加数据行
                for row in rows:
                    data_line = ' & '.join(row)
                    latex_content.append(f'{data_line} \\')
                    latex_content.append('\\hline')
                
                latex_content.append('\\end{tabular}')
                latex_content.append('\\caption{Table}')
                latex_content.append('\\label{tab:example}')
                latex_content.append('\\end{table}')
                latex_content.append('')
        
        # 添加图片
        for image in images:
            image_path = os.path.join('images', image['name'])
            latex_content.append('\\begin{figure}[htbp]')
            latex_content.append('\\centering')
            latex_content.append(f'\\includegraphics[width=0.8\\textwidth]{{{image_path}}}')
            latex_content.append('\\caption{{Image}}')
            latex_content.append(f'\\label{{fig:image_{images.index(image)+1}}}')
            latex_content.append('\\end{figure}')
            latex_content.append('')
        
        # 添加文档尾部
        latex_content.append('\\end{document}')
        
        return '\n'.join(latex_content)
    
    def reconstruct_docx(self, input_dir: str, output_path: str, options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """从拆分素材重新生成DOCX文档
        
        Args:
            input_dir: 输入目录，包含拆分的素材
            output_path: 输出DOCX文件路径
            options: 重构选项
            
        Returns:
            Dict[str, Any]: 重构结果
        """
        if options is None:
            options = {}
        
        # 创建新的DOCX文档
        doc = Document()
        
        # 读取LaTeX文件
        latex_path = os.path.join(input_dir, 'document.tex')
        if not os.path.exists(latex_path):
            return {
                'success': False,
                'error': 'LaTeX file not found'
            }
        
        # 读取LaTeX内容
        with open(latex_path, 'r', encoding='utf-8') as f:
            latex_content = f.read()
        
        # 简化处理，实际项目中需要更复杂的LaTeX解析
        # 这里只是添加一些基本内容作为示例
        doc.add_heading('Reconstructed Document', level=1)
        doc.add_paragraph('This document was reconstructed from split materials.')
        
        # 保存DOCX文档
        doc.save(output_path)
        
        return {
            'success': True,
            'output_path': output_path
        }
    
    def get_supported_formats(self) -> tuple:
        """获取支持的格式"""
        return ('docx', 'markdown')
