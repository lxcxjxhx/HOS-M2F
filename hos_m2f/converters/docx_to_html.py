"""DOCX到HTML格式转换器"""

from typing import Any, Optional, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from hos_m2f.converters.base_converter import BaseConverter
import re


class DOCXToHTMLConverter(BaseConverter):
    """DOCX到HTML格式转换器"""
    
    def convert(self, input_content: bytes, options: Optional[Dict[str, Any]] = None) -> bytes:
        """将DOCX转换为HTML
        
        Args:
            input_content: DOCX文件的二进制数据
            options: 转换选项
                - pretty: 是否美化HTML输出
                - include_styles: 是否包含内联样式
            
        Returns:
            bytes: HTML内容的二进制数据
        """
        if options is None:
            options = {}
        
        pretty = options.get('pretty', True)
        include_styles = options.get('include_styles', True)
        
        # 从二进制数据加载DOCX文档
        import io
        doc = Document(io.BytesIO(input_content))
        
        # 生成HTML内容
        html_content = self._generate_html(doc, pretty, include_styles)
        
        return html_content.encode('utf-8')
    
    def _generate_html(self, doc, pretty: bool, include_styles: bool) -> str:
        """生成HTML内容"""
        html_parts = []
        
        # 添加HTML头部
        html_parts.append('<!DOCTYPE html>')
        html_parts.append('<html>')
        html_parts.append('<head>')
        html_parts.append('    <meta charset="utf-8">')
        html_parts.append('    <title>DOCX to HTML</title>')
        
        # 添加样式
        if include_styles:
            html_parts.append('    <style>')
            html_parts.append(self._generate_css())
            html_parts.append('    </style>')
        
        html_parts.append('</head>')
        html_parts.append('<body>')
        
        # 处理文档内容
        for i, paragraph in enumerate(doc.paragraphs):
            html = self._process_paragraph(paragraph)
            if html:
                if pretty:
                    html_parts.append('    ' + html)
                else:
                    html_parts.append(html)
        
        # 处理表格
        for table in doc.tables:
            html = self._process_table(table)
            if html:
                if pretty:
                    html_parts.append('    ' + html)
                else:
                    html_parts.append(html)
        
        # 添加HTML尾部
        html_parts.append('</body>')
        html_parts.append('</html>')
        
        if pretty:
            return '\n'.join(html_parts)
        else:
            return ''.join(html_parts)
    
    def _generate_css(self) -> str:
        """生成CSS样式"""
        return """
        body {
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            margin: 20px;
        }
        
        h1, h2, h3, h4, h5, h6 {
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            color: #000;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        
        h1 {
            font-size: 24px;
        }
        
        h2 {
            font-size: 20px;
        }
        
        h3 {
            font-size: 16px;
        }
        
        p {
            margin-bottom: 15px;
        }
        
        ul, ol {
            margin-bottom: 15px;
            margin-left: 20px;
        }
        
        li {
            margin-bottom: 5px;
        }
        
        table {
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 20px;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #f2f2f2;
            font-weight: bold;
        }
        
        code {
            font-family: 'Courier New', monospace;
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
        }
        
        pre {
            font-family: 'Courier New', monospace;
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 3px;
            overflow-x: auto;
            margin-bottom: 15px;
        }
        
        a {
            color: #0066cc;
            text-decoration: none;
        }
        
        a:hover {
            text-decoration: underline;
        }
        """
    
    def _process_paragraph(self, paragraph) -> str:
        """处理段落"""
        text = ''
        
        # 检查段落是否为空
        if not paragraph.text.strip():
            return ''
        
        # 检查段落样式
        style_name = paragraph.style.name
        
        # 根据样式确定HTML标签
        if style_name.startswith('Heading 1'):
            tag = 'h1'
        elif style_name.startswith('Heading 2'):
            tag = 'h2'
        elif style_name.startswith('Heading 3'):
            tag = 'h3'
        elif style_name.startswith('Heading 4'):
            tag = 'h4'
        elif style_name.startswith('Heading 5'):
            tag = 'h5'
        elif style_name.startswith('Heading 6'):
            tag = 'h6'
        elif style_name in ['List Bullet', 'List Number']:
            # 列表项由父级处理
            return ''
        else:
            tag = 'p'
        
        # 处理段落中的运行
        for run in paragraph.runs:
            run_html = self._process_run(run)
            text += run_html
        
        # 检查对齐方式
        alignment = paragraph.paragraph_format.alignment
        align_attr = ''
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_attr = ' style="text-align: center;"'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_attr = ' style="text-align: right;"'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_attr = ' style="text-align: justify;"'
        
        # 检查缩进
        indent = paragraph.paragraph_format.left_indent
        if indent and indent.inches > 0:
            if align_attr:
                align_attr = align_attr[:-1] + f'; margin-left: {indent.inches}in;"'
            else:
                align_attr = f' style="margin-left: {indent.inches}in;"'
        
        return f'<{tag}{align_attr}>{text}</{tag}>'
    
    def _process_run(self, run) -> str:
        """处理运行（文本片段）"""
        text = run.text
        if not text:
            return ''
        
        # 处理文本格式
        html_parts = []
        
        # 处理粗体
        if run.bold:
            html_parts.append('<strong>')
        
        # 处理斜体
        if run.italic:
            html_parts.append('<em>')
        
        # 处理下划线
        if run.underline:
            html_parts.append('<u>')
        
        # 处理代码
        if run.font.name == 'Courier New':
            html_parts.append('<code>')
        
        # 添加文本内容（转义HTML特殊字符）
        html_parts.append(self._escape_html(text))
        
        # 处理代码结束标签
        if run.font.name == 'Courier New':
            html_parts.append('</code>')
        
        # 处理下划线结束标签
        if run.underline:
            html_parts.append('</u>')
        
        # 处理斜体结束标签
        if run.italic:
            html_parts.append('</em>')
        
        # 处理粗体结束标签
        if run.bold:
            html_parts.append('</strong>')
        
        return ''.join(html_parts)
    
    def _process_table(self, table) -> str:
        """处理表格"""
        html_parts = []
        html_parts.append('<table>')
        
        # 处理表头
        header_row = table.rows[0]
        html_parts.append('<thead>')
        html_parts.append('<tr>')
        for cell in header_row.cells:
            text = cell.text.strip()
            html_parts.append(f'<th>{self._escape_html(text)}</th>')
        html_parts.append('</tr>')
        html_parts.append('</thead>')
        
        # 处理数据行
        html_parts.append('<tbody>')
        for row in table.rows[1:]:
            html_parts.append('<tr>')
            for cell in row.cells:
                text = cell.text.strip()
                html_parts.append(f'<td>{self._escape_html(text)}</td>')
            html_parts.append('</tr>')
        html_parts.append('</tbody>')
        
        html_parts.append('</table>')
        
        return ''.join(html_parts)
    
    def _escape_html(self, text) -> str:
        """转义HTML特殊字符"""
        escape_map = {
            '&': '&amp;',
            '<': '&lt;',
            '>': '&gt;',
            '"': '&quot;',
            "'": '&#39;'
        }
        
        for char, escape_char in escape_map.items():
            text = text.replace(char, escape_char)
        
        return text
    
    def get_supported_formats(self) -> tuple:
        """获取支持的格式"""
        return ('docx', 'html')