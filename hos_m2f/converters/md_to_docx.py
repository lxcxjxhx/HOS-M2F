"""Markdown到DOCX格式转换器"""

from typing import Any, Optional, Dict
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from hos_m2f.converters.base_converter import BaseConverter
import mistune


class MDToDOCXConverter(BaseConverter):
    """Markdown到DOCX格式转换器"""
    
    def convert(self, input_content: str, options: Optional[Dict[str, Any]] = None) -> bytes:
        """将Markdown转换为DOCX
        
        Args:
            input_content: Markdown内容
            options: 转换选项
                - template: 企业模板文件路径
                - toc: 是否生成目录
                - number_sections: 是否添加章节编号
            
        Returns:
            bytes: DOCX文件的二进制数据
        """
        if options is None:
            options = {}
        
        # 获取模板路径
        template = options.get('template', None)
        
        # 创建文档（使用模板或默认）
        if template and os.path.exists(template):
            doc = Document(template)
        else:
            # 尝试使用默认模板
            try:
                default_template = os.path.join(os.path.dirname(__file__), '..', 'templates', 'reference.docx')
                # 获取绝对路径
                default_template = os.path.abspath(default_template)
                print(f"Trying to use default template: {default_template}")
                if os.path.exists(default_template):
                    doc = Document(default_template)
                else:
                    print(f"Default template not found at: {default_template}")
                    doc = Document()
            except Exception as e:
                print(f"Error loading template: {e}")
                doc = Document()
        
        # 设置默认样式
        self._setup_styles(doc)
        
        # 自定义渲染器
        class DOCXRenderer(mistune.HTMLRenderer):
            def __init__(self, doc):
                super().__init__()
                self.doc = doc
                self.current_paragraph = None
                self.list_level = 0
                self.lists = []
            
            def paragraph(self, text):
                if text.strip():
                    p = self.doc.add_paragraph(text)
                return ''
            
            def heading(self, text, level):
                if level == 1:
                    self.doc.add_heading(text, level=0)
                else:
                    self.doc.add_heading(text, level=level-1)
                return ''
            
            def list(self, text, ordered, start=None, depth=None):
                # 重置列表级别，确保正确的嵌套
                if depth is not None:
                    self.list_level = depth
                else:
                    self.list_level += 1
                self.lists.append(ordered)
                return ''
            
            def list_item(self, text):
                if text.strip():
                    # 使用当前列表级别
                    level = self.list_level
                    if self.lists and level > 0:
                        style = 'List Number' if self.lists[level-1] else 'List Bullet'
                    else:
                        style = 'List Bullet'
                        
                    p = self.doc.add_paragraph(
                        text,
                        style=style
                    )
                    # 缩进
                    if level > 1:
                        # 直接设置缩进，不使用累加
                        p.paragraph_format.left_indent = Inches(0.5 * (level-1))
                return ''
            
            def list_end(self):
                self.list_level -= 1
                if self.lists:
                    self.lists.pop()
                return ''
            
            def table(self, text):
                # 解析Markdown表格并转换为DOCX表格
                try:
                    # 分割表格行
                    rows = text.strip().split('\n')
                    if not rows:
                        return ''
                    
                    # 解析表头
                    header_cells = [cell.strip() for cell in rows[0].split('|') if cell.strip()]
                    if not header_cells:
                        return ''
                    
                    # 创建表格
                    table = self.doc.add_table(rows=1, cols=len(header_cells))
                    table.style = 'Table Grid'
                    
                    # 填充表头
                    header_row = table.rows[0]
                    for i, cell_text in enumerate(header_cells):
                        cell = header_row.cells[i]
                        # 清除单元格默认段落
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.clear()
                        # 添加新的段落和文本
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run(cell_text)
                        run.bold = True
                        run.font.size = Pt(11)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 跳过分隔线行，并解析对齐方式
                    alignments = []
                    if len(rows) > 1 and '---' in rows[1]:
                        # 解析对齐方式
                        alignment_row = rows[1]
                        alignment_cells = [cell.strip() for cell in alignment_row.split('|') if cell.strip()]
                        for cell in alignment_cells:
                            if cell.startswith(':') and cell.endswith(':'):
                                alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                            elif cell.endswith(':'):
                                alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                            else:
                                alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
                        data_rows = rows[2:]
                    else:
                        data_rows = rows[1:]
                    
                    # 填充数据行
                    for row in data_rows:
                        cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                        if cells:
                            new_row = table.add_row()
                            for i, cell_text in enumerate(cells):
                                if i < len(new_row.cells):
                                    cell = new_row.cells[i]
                                    # 清除单元格默认段落
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.clear()
                                    # 添加新的段落和文本
                                    paragraph = cell.add_paragraph()
                                    run = paragraph.add_run(cell_text)
                                    # 设置对齐方式
                                    if i < len(alignments):
                                        paragraph.alignment = alignments[i]
                except Exception as e:
                    # 如果解析失败，回退到简单处理
                    print(f"Error processing table: {e}")
                    self.doc.add_paragraph('Table: ' + text[:100] + '...')
                return ''
            
            def image(self, text, url=None, title=None, alt=None):
                try:
                    # 尝试处理本地和远程图片
                    import os
                    import requests
                    from io import BytesIO
                    
                    # 使用alt作为替代文本
                    if alt is None:
                        alt = text
                    
                    # 检查是否有图片URL
                    if not url:
                        self.doc.add_paragraph(f'Image: {alt}')
                        return ''
                    
                    # 检查是否是本地图片
                    if os.path.exists(url):
                        # 添加本地图片
                        self.doc.add_picture(url)
                    else:
                        # 尝试从网络获取图片
                        response = requests.get(url, timeout=5)
                        if response.status_code == 200:
                            # 添加远程图片
                            image_stream = BytesIO(response.content)
                            self.doc.add_picture(image_stream)
                        else:
                            # 如果获取失败，添加图片描述
                            self.doc.add_paragraph(f'Image: {alt} ({url})')
                except Exception as e:
                    # 如果处理失败，添加图片描述
                    self.doc.add_paragraph(f'Image: {alt or text} ({url or ""})')
                return ''
            
            def link(self, text, url=None, title=None):
                if text and url:
                    p = self.doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    run.underline = WD_UNDERLINE.SINGLE
                    p.add_run(f' ({url})')
                elif text:
                    p = self.doc.add_paragraph(text)
                elif url:
                    p = self.doc.add_paragraph(url)
                return ''
            
            def emphasis(self, text):
                # 直接添加斜体文本
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                return ''
            
            def strong(self, text):
                # 直接添加粗体文本
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                return ''
            
            def codespan(self, text):
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Courier New'
                return ''
            
            def block_code(self, code, info=None):
                # 处理Mermaid图表
                if info == 'mermaid':
                    try:
                        # 尝试渲染Mermaid图表为图片
                        mermaid_image = self._render_mermaid(code)
                        if mermaid_image:
                            # 添加图片
                            self.doc.add_picture(mermaid_image)
                            return ''
                        else:
                            # 如果渲染失败，添加代码块
                            p = self.doc.add_paragraph('Mermaid Chart:')
                            p = self.doc.add_paragraph(code)
                            p.paragraph_format.left_indent = Inches(0.5)
                            return ''
                    except Exception as e:
                        # 如果处理失败，添加代码块
                        p = self.doc.add_paragraph('Mermaid Chart:')
                        p = self.doc.add_paragraph(code)
                        p.paragraph_format.left_indent = Inches(0.5)
                        return ''
                else:
                    # 处理普通代码块
                    p = self.doc.add_paragraph('代码:')
                    p = self.doc.add_paragraph(code)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    return ''
            
            def _render_mermaid(self, mermaid_code):
                """渲染Mermaid图表为图片"""
                # 使用mermaid.ink API渲染Mermaid图表
                try:
                    import requests
                    from io import BytesIO
                    import urllib.parse
                    
                    # 编码Mermaid代码
                    encoded_code = urllib.parse.quote(mermaid_code)
                    
                    # 构建API URL
                    url = f"https://mermaid.ink/img/{encoded_code}"
                    
                    # 发送请求
                    response = requests.get(url, timeout=15, verify=False)
                    
                    if response.status_code == 200:
                        # 检查响应内容是否为图片
                        if response.headers.get('Content-Type', '').startswith('image/'):
                            # 返回图片数据流
                            return BytesIO(response.content)
                        else:
                            # 如果响应不是图片，返回None
                            print(f"Error rendering Mermaid chart: Invalid content type {response.headers.get('Content-Type')}")
                            return None
                    else:
                        # 如果API调用失败，返回None
                        print(f"Error rendering Mermaid chart: HTTP status {response.status_code}")
                        return None
                except Exception as e:
                    # 如果处理失败，返回None
                    print(f"Error rendering Mermaid chart: {e}")
                    return None
        
        # 渲染Markdown
        renderer = DOCXRenderer(doc)
        markdown = mistune.create_markdown(renderer=renderer)
        markdown(input_content)
        
        # 检查是否需要生成目录
        if options.get('toc', False):
            # 在文档开头添加目录
            doc.add_page_break()
            doc.add_heading('目录', level=0)
            doc.add_paragraph().add_run().add_field('TOC')
            doc.add_page_break()
        
        # 保存为二进制数据
        import io
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()
    
    def _setup_styles(self, doc):
        """设置文档样式"""
        styles = doc.styles
        
        # 设置正文样式
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本
        
        # 设置段落格式
        paragraph_format = normal_style.paragraph_format
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.space_after = Pt(12)  # 段后12磅
        
        # 设置标题样式
        for i in range(1, 6):
            heading_style = styles[f'Heading {i}']
            font = heading_style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(14 + (6 - i) * 2)
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本
            
            # 设置标题段落格式
            heading_paragraph_format = heading_style.paragraph_format
            heading_paragraph_format.line_spacing = 1.2  # 1.2倍行距
            heading_paragraph_format.space_after = Pt(12)  # 段后12磅
            heading_paragraph_format.space_before = Pt(0)  # 段前0磅
        
        # 设置列表样式
        for style_name in ['List Bullet', 'List Number']:
            if style_name in styles:
                list_style = styles[style_name]
                font = list_style.font
                font.name = 'Microsoft YaHei'
                font.size = Pt(12)
                font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本
                
                # 设置列表段落格式
                list_paragraph_format = list_style.paragraph_format
                list_paragraph_format.line_spacing = 1.5  # 1.5倍行距
                list_paragraph_format.space_after = Pt(6)  # 段后6磅
    
    def get_supported_formats(self) -> tuple:
        """获取支持的格式"""
        return ('markdown', 'docx')
